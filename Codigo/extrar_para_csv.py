# ================================================================
# SCRIPT PARA EXTRAIR DADOS DE DOCX E GERAR CSVs ESTRUTURADOS
# ================================================================

import re
import pandas as pd
from docx import Document
import os

# ================================================================
# 1. CONFIGURAÇÃO
# ================================================================

# Nomes dos arquivos DOCX
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_ENTREVISTAS = os.path.join(BASE_DIR, "Trechos_Entrevistas.docx")
DOCX_MOBILIZACAO = os.path.join(BASE_DIR, "Registros_Mobilizacao.docx")
DOCX_INDICADORES = os.path.join(BASE_DIR, "Indicadores_Secundarios.docx")

# Nomes dos CSVs de saída
CSV_ENTREVISTAS = "entrevistas.csv"
CSV_MOBILIZACAO = "mobilizacao.csv"
CSV_INDICADORES = "indicadores_sec.csv"

# ================================================================
# 2. FUNÇÕES DE EXTRAÇÃO
# ================================================================

def extrair_texto_docx(caminho):
    """Extrai todo o texto de um arquivo DOCX."""
    doc = Document(caminho)
    texto_completo = ""
    for paragrafo in doc.paragraphs:
        if paragrafo.text.strip():
            texto_completo += paragrafo.text + "\n"
    return texto_completo

def extrair_entrevistas(docx_path):
    """Extrai dados de entrevistas. Saída: id, perfil, municipio, zona, trecho"""
    dados = []
    texto_completo = extrair_texto_docx(docx_path)
    
    padrao = r"Entrevista\s+(\d+)\s*[–\-]\s*([^,]+),\s*(\d+)\s*anos,\s*([^()]+)\s*\(([^)]+)\)"
    matches = list(re.finditer(padrao, texto_completo, re.IGNORECASE))
    
    for i, match in enumerate(matches):
        id_ent = int(match.group(1))
        perfil = match.group(2).strip()
        idade = match.group(3).strip()
        municipio = match.group(4).strip()
        zona = match.group(5).strip()
        perfil_completo = f"{perfil}, {idade} anos"
        
        start = match.end()
        end = matches[i+1].start() if i+1 < len(matches) else len(texto_completo)
        trecho = texto_completo[start:end].strip()
        trecho = re.sub(r"Trecho:", "", trecho, flags=re.IGNORECASE).strip()
        trecho = " ".join(trecho.split())
        
        dados.append({
            "id": id_ent,
            "perfil": perfil_completo,
            "municipio": municipio,
            "zona": zona,
            "trecho": trecho
        })
    
    return pd.DataFrame(dados)

def extrair_mobilizacao(docx_path):
    """Extrai registros de mobilização. Saída: id, municipio, data, resumo, sentimento"""
    dados = []
    texto_completo = extrair_texto_docx(docx_path)
    
    padrao = r"Registro\s+(\d+)\s*[–\-]\s*([^\n]+)"
    matches = list(re.finditer(padrao, texto_completo, re.IGNORECASE))
    
    for i, match in enumerate(matches):
        id_reg = int(match.group(1))
        municipio = match.group(2).strip()
        
        start = match.end()
        end = matches[i+1].start() if i+1 < len(matches) else len(texto_completo)
        bloco = texto_completo[start:end].strip()
        
        data_match = re.search(r"Data:\s*(\d{2}/\d{2}/\d{4})", bloco, re.IGNORECASE)
        data = data_match.group(1) if data_match else ""
        
        resumo_match = re.search(r"Resumo:\s*(.+)", bloco, re.DOTALL | re.IGNORECASE)
        resumo = resumo_match.group(1).strip() if resumo_match else bloco
        resumo = " ".join(resumo.split())
        
        # Classificar sentimento
        positivo = ["boa", "receptividade", "confiança", "entender", "participação", "interesse", "incentivar", "fortaleceu", "adesão", "compartilharam"]
        negativo = ["baixa", "resistência", "dúvidas", "interrupção", "escassez", "reclamação", "receio", "preocupações"]
        score = 0
        for p in positivo:
            if p in resumo.lower():
                score += 1
        for n in negativo:
            if n in resumo.lower():
                score -= 1
        sentimento = "positivo" if score > 0 else "negativo" if score < 0 else "neutro"
        
        dados.append({
            "id": id_reg,
            "municipio": municipio,
            "data": data,
            "resumo": resumo,
            "sentimento": sentimento
        })
    
    return pd.DataFrame(dados)

def extrair_indicadores(docx_path):
    """
    Extrai a tabela de indicadores secundários do DOCX.
    Saída: DataFrame com colunas:
        Municipio, % Populacao Rural, Cobertura APS (%), Desocupacao (%),
        Jovens 18–29 (%), Domicilios com agua encanada (%), Observacao
    """
    doc = Document(docx_path)
    dados = []

    # 1. Tentar extrair de tabelas do DOCX
    for table in doc.tables:
        # Verifica se a tabela tem cabeçalho com os nomes esperados
        cabecalho = [cell.text.strip().lower() for cell in table.rows[0].cells]
        
        # Se encontrar palavras-chave, considera que é a tabela de indicadores
        palavras_chave = ["município", "rural", "cobertura", "desocupação", "jovens", "água"]
        if any(p in " ".join(cabecalho) for p in palavras_chave):
            # Itera pelas linhas de dados (a partir da linha 1)
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                # Garante que há células suficientes
                if len(cells) >= 7:
                    dados.append({
                        "Municipio": cells[0],
                        "% Populacao Rural": cells[1] if len(cells) > 1 else "",
                        "Cobertura APS (%)": cells[2] if len(cells) > 2 else "",
                        "Desocupacao (%)": cells[3] if len(cells) > 3 else "",
                        "Jovens 18–29 (%)": cells[4] if len(cells) > 4 else "",
                        "Domicilios com agua encanada (%)": cells[5] if len(cells) > 5 else "",
                        "Observacao": cells[6] if len(cells) > 6 else ""
                    })
            break  # encontrou a tabela, sai do loop

    # 2. Se não encontrou tabela, tenta extrair do texto (fallback)
    if not dados:
        texto_completo = extrair_texto_docx(docx_path)
        linhas = texto_completo.split("\n")
        municipio_atual = None
        dados_linha = {}
        
        for linha in linhas:
            linha = linha.strip()
            if not linha:
                continue
            
            # Tentar identificar municípios
            municipios_conhecidos = ["Rio Claro", "Santa Aurora", "Vale Verde", "Serra Azul", "Boa Esperança", "Lagoa Nova"]
            for mun in municipios_conhecidos:
                if mun in linha and len(linha) < 50:
                    # Se já havia dados do município anterior, salvar
                    if municipio_atual and dados_linha:
                        dados.append({
                            "Municipio": municipio_atual,
                            "% Populacao Rural": dados_linha.get("pop_rural", ""),
                            "Cobertura APS (%)": dados_linha.get("cobertura_aps", ""),
                            "Desocupacao (%)": dados_linha.get("desocupacao", ""),
                            "Jovens 18–29 (%)": dados_linha.get("jovens", ""),
                            "Domicilios com agua encanada (%)": dados_linha.get("agua", ""),
                            "Observacao": dados_linha.get("obs", "")
                        })
                    municipio_atual = mun
                    dados_linha = {}
                    break
            
            # Extrair valores numéricos
            padrao_valor = r"([\d,]+)\s*%?"
            if "população rural" in linha.lower():
                match = re.search(padrao_valor, linha)
                if match:
                    dados_linha["pop_rural"] = match.group(1).replace(",", ".")
            elif "cobertura aps" in linha.lower():
                match = re.search(padrao_valor, linha)
                if match:
                    dados_linha["cobertura_aps"] = match.group(1).replace(",", ".")
            elif "desocupação" in linha.lower():
                match = re.search(padrao_valor, linha)
                if match:
                    dados_linha["desocupacao"] = match.group(1).replace(",", ".")
            elif "jovens" in linha.lower():
                match = re.search(padrao_valor, linha)
                if match:
                    dados_linha["jovens"] = match.group(1).replace(",", ".")
            elif "água" in linha.lower() or "domicílios" in linha.lower():
                match = re.search(padrao_valor, linha)
                if match:
                    dados_linha["agua"] = match.group(1).replace(",", ".")
        
        # Salvar último município
        if municipio_atual and dados_linha:
            dados.append({
                "Municipio": municipio_atual,
                "% Populacao Rural": dados_linha.get("pop_rural", ""),
                "Cobertura APS (%)": dados_linha.get("cobertura_aps", ""),
                "Desocupacao (%)": dados_linha.get("desocupacao", ""),
                "Jovens 18–29 (%)": dados_linha.get("jovens", ""),
                "Domicilios com agua encanada (%)": dados_linha.get("agua", ""),
                "Observacao": dados_linha.get("obs", "")
            })

    return pd.DataFrame(dados)

# ================================================================
# 3. EXECUÇÃO PRINCIPAL
# ================================================================

def main():
    print("="*70)
    print("EXTRAÇÃO DE DADOS DE DOCX PARA CSV")
    print("="*70)
    print(f"📂 Diretório base: {BASE_DIR}")
    
    arquivos = {
        "entrevistas": DOCX_ENTREVISTAS,
        "mobilizacao": DOCX_MOBILIZACAO,
        "indicadores": DOCX_INDICADORES
    }
    
    for nome, caminho in arquivos.items():
        if not os.path.exists(caminho):
            print(f"⚠️ Arquivo não encontrado: {caminho}")
            arquivos[nome] = None
    
    # Entrevistas
    if arquivos["entrevistas"]:
        print(f"\n📄 Extraindo entrevistas...")
        df_entrev = extrair_entrevistas(DOCX_ENTREVISTAS)
        df_entrev.to_csv(CSV_ENTREVISTAS, index=False, encoding='utf-8-sig')
        print(f"✅ {CSV_ENTREVISTAS} ({len(df_entrev)} registros)")
        print(df_entrev[['id', 'perfil', 'municipio', 'zona']].head())
    else:
        print("❌ Nenhum arquivo de entrevistas encontrado.")
    
    # Mobilização
    if arquivos["mobilizacao"]:
        print(f"\n📄 Extraindo mobilização...")
        df_mob = extrair_mobilizacao(DOCX_MOBILIZACAO)
        df_mob.to_csv(CSV_MOBILIZACAO, index=False, encoding='utf-8-sig')
        print(f"✅ {CSV_MOBILIZACAO} ({len(df_mob)} registros)")
        print(df_mob[['id', 'municipio', 'data', 'sentimento']].head())
    else:
        print("❌ Nenhum arquivo de mobilização encontrado.")
    
    # Indicadores
    if arquivos["indicadores"]:
        print(f"\n📄 Extraindo indicadores secundários...")
        df_ind = extrair_indicadores(DOCX_INDICADORES)
        if not df_ind.empty:
            df_ind.to_csv(CSV_INDICADORES, index=False, encoding='utf-8-sig')
            print(f"Colunas: {list(df_ind.columns)}")
            print(df_ind.head())
        else:
            print(f"⚠️ Nenhum indicador encontrado. {CSV_INDICADORES} não foi gerado.")
    else:
        print("❌ Nenhum arquivo de indicadores encontrado.")
    
    print("\n" + "="*70)
    print("✅ EXTRAÇÃO CONCLUÍDA!")
    print("="*70)

if __name__ == "__main__":
    main()